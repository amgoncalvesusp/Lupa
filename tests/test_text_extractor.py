"""Unit tests for PDF/DOCX/TXT text extraction."""

import pytest
from docx import Document

from src.core.pdf_processor import DocumentProcessor, PDFProcessor
from src.core.text_extractor import extract_pages

pytestmark = pytest.mark.unit


def test_docx_extracts_text(tmp_path):
    path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("Primeiro parágrafo")
    doc.add_paragraph("Segundo parágrafo")
    doc.save(path)

    pages, stats = extract_pages(path)

    assert "Primeiro parágrafo" in pages[0]
    assert "Segundo parágrafo" in pages[0]
    assert stats["ocr_pages_count"] == 0


def test_txt_extracts_two_blocks(tmp_path):
    path = tmp_path / "doc.txt"
    path.write_text("Bloco um\n\nBloco dois", encoding="utf-8")

    pages, stats = extract_pages(path)

    assert pages == ["Bloco um", "Bloco dois"]
    assert stats["pages_with_text"] == 2


def test_unknown_extension_raises_clear_value_error(tmp_path):
    path = tmp_path / "doc.md"
    path.write_text("x", encoding="utf-8")

    with pytest.raises(ValueError, match="Formato não suportado"):
        extract_pages(path)


def test_pdfprocessor_alias_points_to_document_processor():
    assert PDFProcessor is DocumentProcessor


def test_document_processor_handles_txt_end_to_end(tmp_path):
    path = tmp_path / "doc.txt"
    path.write_text("Este documento fala de clima e território.", encoding="utf-8")

    result = DocumentProcessor(
        enable_ocr=False,
        search_terms=[("clima", False)],
        detect_sentiment=False,
        detect_emotions=False,
        detect_textmetrics=False,
        detect_kwic=False,
    ).process(str(path))

    assert result["filename"] == "doc.txt"
    assert result["words_total"] >= 6
    assert result["term_results"]["clima"]["analytical"] == 1
